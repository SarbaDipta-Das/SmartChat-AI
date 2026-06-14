import os
import re
from pathlib import Path
from typing import Optional
import docx
import pdfplumber
import fitz  # PyMuPDF
import faiss
import numpy as np
import pickle

# OCR support
try:
    import pytesseract
    from PIL import Image
    import io
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

INDEX_PATH = "faiss_index.bin"
CHUNKS_PATH = "chunks.pkl"

MAX_PAGES = 10  # maximum pages to process

_embedder = None

def get_embedder():
    global _embedder
    if _embedder is None:
        from sentence_transformers import SentenceTransformer
        _embedder = SentenceTransformer("all-MiniLM-L6-v2")
    return _embedder


def extract_text_from_pdf(path: str) -> str:
    """Smart PDF extraction — tries 3 methods automatically."""
    all_text = []

    # ── Method 1: pdfplumber (best for normal text PDFs) ─────────────────────
    try:
        with pdfplumber.open(path) as pdf:
            pages = pdf.pages[:MAX_PAGES]
            for i, page in enumerate(pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    all_text.append(f"[Page {i+1}]\n{page_text.strip()}")
    except Exception as e:
        print(f"pdfplumber error: {e}")

    text_so_far = "\n\n".join(all_text)
    if len(text_so_far.strip()) > 200:
        print(f"✅ pdfplumber extracted {len(text_so_far)} chars from {path}")
        return text_so_far

    # ── Method 2: PyMuPDF (good for complex PDFs) ────────────────────────────
    all_text = []
    try:
        doc = fitz.open(path)
        pages_to_read = min(len(doc), MAX_PAGES)
        for i in range(pages_to_read):
            page = doc[i]
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                all_text.append(f"[Page {i+1}]\n{page_text.strip()}")
        doc.close()
    except Exception as e:
        print(f"PyMuPDF error: {e}")

    text_so_far = "\n\n".join(all_text)
    if len(text_so_far.strip()) > 200:
        print(f"✅ PyMuPDF extracted {len(text_so_far)} chars from {path}")
        return text_so_far

    # ── Method 3: OCR (for scanned/image PDFs) ───────────────────────────────
    if OCR_AVAILABLE:
        print("📷 Using OCR for scanned PDF...")
        all_text = []
        try:
            doc = fitz.open(path)
            pages_to_read = min(len(doc), MAX_PAGES)
            for i in range(pages_to_read):
                page = doc[i]
                # render at 3x zoom for better OCR accuracy
                mat = fitz.Matrix(3, 3)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                # OCR config for better accuracy
                custom_config = r'--oem 3 --psm 6'
                page_text = pytesseract.image_to_string(img, lang='eng', config=custom_config)
                if page_text and page_text.strip():
                    all_text.append(f"[Page {i+1}]\n{page_text.strip()}")
                print(f"  OCR page {i+1}/{pages_to_read} done")
            doc.close()
        except Exception as e:
            print(f"OCR error: {e}")

        ocr_text = "\n\n".join(all_text)
        if ocr_text.strip():
            print(f"✅ OCR extracted {len(ocr_text)} chars from {path}")
            return ocr_text
    else:
        print("⚠️ OCR not available. Install pytesseract for scanned PDFs.")

    # Return whatever we have
    final = "\n\n".join(all_text)
    print(f"⚠️ Final extracted text: {len(final)} chars")
    return final


def extract_text_from_docx(path: str) -> str:
    """Extract text from DOCX including tables."""
    doc = docx.Document(path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(path)
    else:
        return extract_text_from_txt(path)


def clean_text(text: str) -> str:
    """Clean extracted text — remove garbage characters."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {3,}', ' ', text)
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
    # Remove very short lines (likely OCR noise)
    lines = [line for line in text.split('\n') if len(line.strip()) > 3]
    return '\n'.join(lines)


def chunk_text(text: str, chunk_size: int = 400, overlap: int = 80) -> list:
    """Smart chunking — splits by sentences with overlap."""
    text = clean_text(text)
    # Split by sentence endings
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks, current, count = [], [], 0

    for sentence in sentences:
        words = sentence.split()
        if not words:
            continue
        if count + len(words) > chunk_size and current:
            chunk_text_str = " ".join(current).strip()
            if len(chunk_text_str) > 50:  # only keep meaningful chunks
                chunks.append(chunk_text_str)
            # keep overlap for context
            overlap_words = current[-overlap:] if overlap < len(current) else current[:]
            current = overlap_words + words
            count = len(current)
        else:
            current.extend(words)
            count += len(words)

    if current:
        chunk_text_str = " ".join(current).strip()
        if len(chunk_text_str) > 50:
            chunks.append(chunk_text_str)

    return chunks


def load_index():
    if os.path.exists(INDEX_PATH) and os.path.exists(CHUNKS_PATH):
        index = faiss.read_index(INDEX_PATH)
        with open(CHUNKS_PATH, "rb") as f:
            chunks = pickle.load(f)
        return index, chunks
    return None, []


def save_index(index, chunks):
    faiss.write_index(index, INDEX_PATH)
    with open(CHUNKS_PATH, "wb") as f:
        pickle.dump(chunks, f)


def add_document_to_index(file_path: str) -> int:
    """Extract, chunk and index a document."""
    print(f"📄 Processing: {file_path}")
    text = extract_text(file_path)

    if not text or len(text.strip()) < 50:
        print(f"❌ No text extracted from {file_path}")
        return 0

    print(f"📝 Total text: {len(text)} characters")
    new_chunks = chunk_text(text)
    print(f"🔪 Created {len(new_chunks)} chunks")

    if not new_chunks:
        return 0

    embedder = get_embedder()
    embeddings = embedder.encode(new_chunks, show_progress_bar=False)
    embeddings = np.array(embeddings).astype("float32")

    index, existing_chunks = load_index()
    if index is None:
        dim = embeddings.shape[1]
        index = faiss.IndexFlatL2(dim)

    index.add(embeddings)
    existing_chunks.extend(new_chunks)
    save_index(index, existing_chunks)
    print(f"✅ Successfully indexed {len(new_chunks)} chunks!")
    return len(new_chunks)


def retrieve_context(query: str, top_k: int = 5) -> Optional[str]:
    """Retrieve most relevant chunks for a query."""
    index, chunks = load_index()
    if index is None or not chunks:
        return None

    embedder = get_embedder()
    query_vec = embedder.encode([query]).astype("float32")
    distances, indices = index.search(query_vec, min(top_k, len(chunks)))

    results = []
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(chunks) and dist < 2.5:  # slightly relaxed threshold
            results.append(chunks[idx])

    return "\n\n".join(results) if results else None


def clear_index():
    for path in [INDEX_PATH, CHUNKS_PATH]:
        if os.path.exists(path):
            os.remove(path)